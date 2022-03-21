import io
import pathlib
import subprocess
from itertools import chain
from logging import Logger
from tempfile import gettempdir
from typing import Any, Callable, Dict, List

from docx import Document
from template_system.application import protocols
from template_system.application.protocols import TemplateDataProtocol
from template_system.domain import TemplateData
from template_system.errors import (
    IncompatibleTemplateAndExporterError,
    TemplateProcessError,
)


class DocxTemplateVarsApplier:
    def __init__(self, template_factory: Callable = Document) -> None:
        self.__template_items: List[Any] = []
        self.__temp_dir = gettempdir()
        self.__template_factory = template_factory

    def apply_template_vars(self, template_data: TemplateDataProtocol) -> Any:
        template = self.__template_factory(io.BytesIO(template_data.input))
        self.__process_section(template=template)
        self.__process_tables(template=template)
        self.__process_paragraphs(template=template)

        output_document_path = pathlib.Path(
            self.__temp_dir,
            f"{template_data.unique_identifier}.docx",
        )
        for key, value in template_data.template_vars.items():
            for item in self.__template_items:
                if key in item.text:
                    item.text = item.text.replace(key, value)

        template.save(output_document_path)
        template_data.input = output_document_path
        return template_data

    def __process_section(self, template: Any):
        for section in template.sections:
            self.__template_items += section.header.paragraphs
            self.__template_items += section.footer.paragraphs
            self.__process_tables(
                tables=section.header.tables,
                template=template,
            )
            self.__process_tables(
                tables=section.footer.tables,
                template=template,
            )

    def __process_tables(
        self,
        *,
        template: Any,
        tables=None,
    ) -> None:
        tables = template.tables if not tables else tables
        cells = []

        for table in tables:
            for item in chain(table.columns, table.rows):
                cells += item.cells

        for cell in cells:
            self.__template_items += cell.paragraphs
            if cell.tables:
                self.__process_tables(tables=cell.tables, template=template)

    def __process_paragraphs(self, template) -> None:
        self.__template_items += template.paragraphs


VarsApplier = DocxTemplateVarsApplier


class LibreOfficeDocxToPDFTemplateExporter(protocols.TemplateExporterABC):
    suported_types = ("docx",)
    output_type = "pdf"

    def __init__(
        self,
        template_vars_applier: protocols.TemplateApplier = VarsApplier(),
        subprocess_executor: Callable = subprocess.run,
    ) -> None:
        self._template_vars_applier = template_vars_applier
        self._subprocess_executor = subprocess_executor
        super().__init__()

    def __export_pdf(self, input_file: str, output_dir=gettempdir()) -> str:
        filename_with_extension = pathlib.Path(input_file).name
        filename_without_extension = filename_with_extension.replace(
            f".{self.suported_types[0]}", ""
        )
        output_file = pathlib.Path(
            output_dir, f"{filename_without_extension}.{self.output_type}"
        )

        self._subprocess_executor(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                input_file,
                "--outdir",
                output_dir,
            ]
        )
        return str(output_file)

    def _validate_template_data(
        self, template_data: TemplateData
    ) -> TemplateDataProtocol:
        # TODO: create tests to the input_file types -> str | path| bytes

        error_msg = "Unsuported template input"

        if not template_data.input:
            raise IncompatibleTemplateAndExporterError(error_msg)

        if type(template_data.input) in (str, pathlib.Path):
            template_data.input = str(template_data.input)
            if template_data.input.split(".")[-1].lower() != self.suported_types[0]:
                raise IncompatibleTemplateAndExporterError(error_msg)

    def config_template_data(
        self,
        template_data: TemplateData,
    ) -> TemplateDataProtocol:
        # TODO: create tests to the input_file types -> str | path| bytes

        self._validate_template_data(template_data)

        return TemplateData(
            input=template_data.input,  # type: ignore
            input_type=self.suported_types[0],
            output="",
            output_type=self.output_type,
            template_vars=template_data.template_vars,
        )

    def process(
        self, template_data: protocols.TemplateDataProtocol
    ) -> protocols.TemplateDataProtocol:
        assert type(template_data.input) is bytes

        try:
            self._template_vars_applier.apply_template_vars(template_data)
            template_data.output = self.__export_pdf(template_data.input)
            return template_data
        except Exception as err:
            error_msg = "Error on processing template"
            raise TemplateProcessError(error_msg) from err

    def __str__(self) -> str:
        return f"{self.__class__.__name__}"

    def __repr__(self) -> str:
        return self.__str__()
